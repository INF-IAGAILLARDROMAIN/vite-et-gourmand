import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# === EN-TETE ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EVALUATION EN COURS DE FORMATION')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x78, 0x35, 0x0F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TP \u2013 D\u00e9veloppeur Web et Web Mobile')
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph()

# Info table
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
info = [
    ('NOM', 'GAILLARD'),
    ('Pr\u00e9nom', 'Romain'),
    ('Date de naissance', '12/06/1998'),
    ('Lieu de naissance', 'Auray (56 - Morbihan)'),
]
for i, (label, value) in enumerate(info):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = value

doc.add_paragraph()

# === LIENS OBLIGATOIRES ===
h = doc.add_heading('Liens obligatoires', level=1)
for r in h.runs:
    r.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

p = doc.add_paragraph()
run = p.add_run('SANS CES ELEMENTS, VOTRE COPIE SERA REJETEE')
run.bold = True
run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
run.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

links = [
    ('Lien du Git', 'https://github.com/INF-IAGAILLARDROMAIN/vite-et-gourmand'),
    ('Lien de l\u2019outil de gestion de projet', 'https://github.com/INF-IAGAILLARDROMAIN/vite-et-gourmand/issues'),
    ('Lien du d\u00e9ploiement', 'https://vite-et-gourmand-rust.vercel.app'),
    ('Login et mot de passe administrateur', 'admin@viteetgourmand.fr / Admin@2026!'),
]
for label, value in links:
    p = doc.add_paragraph()
    run = p.add_run(label + ' : ')
    run.bold = True
    p.add_run(value)

doc.add_page_break()

# === PARTIE 1 ===
h = doc.add_heading('Partie 1 : Analyse des besoins', level=1)
for r in h.runs:
    r.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

doc.add_heading('1.1 R\u00e9sum\u00e9 du projet (200-250 mots)', level=2)

resume = (
    "Vite & Gourmand est une application web d\u00e9velopp\u00e9e pour un traiteur "
    "\u00e9v\u00e9nementiel bas\u00e9 \u00e0 Bordeaux. L\u2019entreprise, dirig\u00e9e par Jos\u00e9, propose des "
    "menus gastronomiques livr\u00e9s \u00e0 domicile pour tous types d\u2019\u00e9v\u00e9nements : "
    "mariages, anniversaires, s\u00e9minaires d\u2019entreprise, et f\u00eates de famille.\n\n"
    "L\u2019application a pour objectif de digitaliser l\u2019ensemble du processus "
    "m\u00e9tier du traiteur : de la pr\u00e9sentation du catalogue de menus \u00e0 la "
    "gestion compl\u00e8te des commandes, en passant par le suivi de livraison "
    "et la collecte d\u2019avis clients.\n\n"
    "Le syst\u00e8me s\u2019articule autour de trois profils d\u2019utilisateurs. Les "
    "visiteurs et clients peuvent consulter les menus, filtrer par th\u00e8me "
    "(No\u00ebl, P\u00e2ques, Classique, \u00c9v\u00e9nement) ou r\u00e9gime alimentaire "
    "(v\u00e9g\u00e9tarien, v\u00e9gan, sans gluten), passer commande avec calcul "
    "automatique du prix (incluant frais de livraison et r\u00e9ductions "
    "\u00e9ventuelles), et suivre l\u2019avancement de leur commande en temps r\u00e9el.\n\n"
    "Les employ\u00e9s disposent d\u2019un back-office pour g\u00e9rer les commandes "
    "selon un workflow pr\u00e9cis (de \u00abRe\u00e7ue\u00bb \u00e0 \u00abTermin\u00e9e\u00bb), mod\u00e9rer les "
    "avis clients, et administrer les menus et horaires.\n\n"
    "L\u2019administrateur b\u00e9n\u00e9ficie de fonctionnalit\u00e9s suppl\u00e9mentaires : "
    "cr\u00e9ation et d\u00e9sactivation de comptes employ\u00e9s, et consultation de "
    "statistiques de performance (commandes par menu, chiffre d\u2019affaires) "
    "via des graphiques aliment\u00e9s par une base de donn\u00e9es NoSQL (MongoDB).\n\n"
    "L\u2019application int\u00e8gre \u00e9galement un syst\u00e8me complet d\u2019emails "
    "automatiques (confirmation de commande, relance retour mat\u00e9riel, "
    "invitation \u00e0 laisser un avis) et respecte les normes de s\u00e9curit\u00e9 "
    "(RGPD, accessibilit\u00e9 RGAA)."
)
doc.add_paragraph(resume)

doc.add_heading('1.2 Cahier des charges / Sp\u00e9cifications fonctionnelles', level=2)

doc.add_paragraph(
    'Contexte : Le traiteur Vite & Gourmand souhaite moderniser sa '
    'pr\u00e9sence en ligne et automatiser la gestion de ses commandes '
    '\u00e9v\u00e9nementielles.'
)

doc.add_paragraph('Objectifs :')
objectives = [
    'Permettre aux clients de consulter et commander des menus en ligne',
    'Automatiser le calcul de prix (menu + livraison + r\u00e9ductions)',
    'Fournir un outil de gestion des commandes pour l\u2019\u00e9quipe',
    'Offrir un tableau de bord statistique \u00e0 l\u2019administrateur',
    'Garantir la s\u00e9curit\u00e9 des donn\u00e9es et la conformit\u00e9 RGPD',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Number')

doc.add_paragraph('Fonctionnalit\u00e9s principales :')

table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Module'
table.rows[0].cells[1].text = 'Fonctionnalit\u00e9s'
for cell in table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

modules = [
    ('Catalogue', 'Affichage des menus avec filtres (th\u00e8me, r\u00e9gime, budget), d\u00e9tail avec composition et allerg\u00e8nes'),
    ('Commande', 'Cr\u00e9ation avec calcul prix dynamique, suivi par statut, annulation, historique complet'),
    ('Authentification', 'Inscription s\u00e9curis\u00e9e, connexion JWT, r\u00e9initialisation mot de passe'),
    ('Avis', 'D\u00e9p\u00f4t d\u2019avis (note + commentaire) sur commandes termin\u00e9es, mod\u00e9ration par l\u2019\u00e9quipe'),
    ('Back-office', 'Gestion commandes (workflow 7 statuts), menus, horaires, avis'),
    ('Administration', 'Cr\u00e9ation/d\u00e9sactivation employ\u00e9s, statistiques MongoDB (graphiques Recharts)'),
    ('Emails', '7 templates automatiques (bienvenue, confirmation, retour mat\u00e9riel, etc.)'),
    ('Contact', 'Formulaire public avec envoi d\u2019email \u00e0 l\u2019entreprise'),
]
for i, (module, desc) in enumerate(modules):
    table.rows[i+1].cells[0].text = module
    table.rows[i+1].cells[1].text = desc

doc.add_page_break()

# === PARTIE 2 ===
h = doc.add_heading('Partie 2 : Sp\u00e9cifications techniques', level=1)
for r in h.runs:
    r.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

doc.add_heading('2.1 Technologies utilis\u00e9es et justification', level=2)

table = doc.add_table(rows=15, cols=3)
table.style = 'Table Grid'
for j, h_text in enumerate(['Technologie', 'R\u00f4le', 'Justification']):
    table.rows[0].cells[j].text = h_text
    table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

techs = [
    ('Next.js 16', 'Frontend', 'Framework React avec SSR natif pour le SEO, App Router, optimisations de performance automatiques'),
    ('React 19', 'UI', 'Derni\u00e8re version stable, hooks pour la gestion d\u2019\u00e9tat, large \u00e9cosyst\u00e8me'),
    ('TypeScript 5', 'Typage', 'Typage statique pour r\u00e9duire les bugs, autocompl\u00e9tion IDE'),
    ('Tailwind CSS v4', 'Styles', 'Approche utility-first, bundle CSS minimal gr\u00e2ce au purge automatique'),
    ('Framer Motion', 'Animations', 'Animations GPU-acc\u00e9l\u00e9r\u00e9es, API d\u00e9clarative, int\u00e9gration React'),
    ('NestJS 11', 'Backend', 'Architecture modulaire avec injection de d\u00e9pendances, support TypeScript natif'),
    ('Prisma 7', 'ORM SQL', 'G\u00e9n\u00e9ration automatique des types TypeScript, migrations simplifi\u00e9es'),
    ('PostgreSQL 16', 'BDD SQL', 'Base relationnelle robuste, conforme ACID, h\u00e9bergement Neon'),
    ('MongoDB Atlas', 'BDD NoSQL', 'Stockage flexible pour les statistiques, pipelines d\u2019agr\u00e9gation'),
    ('Mongoose', 'ODM NoSQL', 'Int\u00e9gration NestJS native, sch\u00e9mas TypeScript'),
    ('JWT + Passport', 'Auth', 'Tokens stateless, strat\u00e9gie \u00e9prouv\u00e9e, int\u00e9gration NestJS native'),
    ('Bcrypt', 'S\u00e9curit\u00e9', 'Hachage de mots de passe avec salt, 10 rounds'),
    ('Nodemailer', 'Emails', 'Envoi SMTP universel, templates HTML'),
    ('Recharts', 'Graphiques', 'Biblioth\u00e8que React pour les graphiques, composants d\u00e9claratifs'),
]
for i, (tech, role, justif) in enumerate(techs):
    table.rows[i+1].cells[0].text = tech
    table.rows[i+1].cells[1].text = role
    table.rows[i+1].cells[2].text = justif

doc.add_heading('2.2 Mise en place de l\u2019environnement de travail', level=2)

doc.add_paragraph(
    "L\u2019environnement de travail est document\u00e9 dans le fichier README.md du "
    "repository GitHub. Structure monorepo avec apps/web (Next.js) et apps/api (NestJS).\n\n"
    "Pr\u00e9requis : Node.js >= 20, npm >= 10, PostgreSQL (ou Neon), MongoDB (ou Atlas).\n\n"
    "Installation :\n"
    "1. Clone du repo : git clone https://github.com/INF-IAGAILLARDROMAIN/vite-et-gourmand.git\n"
    "2. Backend : cd apps/api && npm install\n"
    "3. Frontend : cd apps/web && npm install\n"
    "4. Configuration : fichiers .env (API) et .env.local (Web)\n"
    "5. Initialisation BDD : npx prisma generate && npx prisma migrate dev && npx prisma db seed\n"
    "6. D\u00e9marrage : API sur port 3000, Web sur port 3001\n\n"
    "Workflow Git : branches main (production) \u2192 develop (int\u00e9gration) \u2192 feature/* (d\u00e9veloppement). "
    "Convention de commits s\u00e9mantiques (feat, fix, docs, chore, test)."
)

doc.add_heading('2.3 M\u00e9canismes de s\u00e9curit\u00e9', level=2)

sec_items = [
    'Validation c\u00f4t\u00e9 client ET serveur : tous les DTOs NestJS utilisent class-validator (@IsEmail(), @MinLength(), @Matches())',
    'Politique de mot de passe fort : minimum 10 caract\u00e8res, 1 majuscule, 1 minuscule, 1 chiffre, 1 caract\u00e8re sp\u00e9cial',
    'JWT : tokens sign\u00e9s avec un secret configurable, expiration param\u00e9trable (7 jours par d\u00e9faut)',
    'Bcrypt : hachage des mots de passe avec salt (10 rounds), mots de passe jamais stock\u00e9s en clair',
    'RBAC : contr\u00f4le d\u2019acc\u00e8s par r\u00f4les via @Roles() decorator et RolesGuard NestJS',
    'CORS : origines autoris\u00e9es limit\u00e9es au domaine frontend uniquement',
    'Anti-\u00e9num\u00e9ration : la route forgot-password retourne toujours un succ\u00e8s, m\u00eame si l\u2019email n\u2019existe pas',
    'Validation stricte : forbidNonWhitelisted: true sur le ValidationPipe global',
    'Protection XSS : React \u00e9chappe automatiquement les variables dans le JSX',
    'Protection CSRF : utilisation de JWT dans les headers Authorization (pas de cookies de session)',
    'Conformit\u00e9 RGPD : mentions l\u00e9gales, droit d\u2019acc\u00e8s et suppression, mots de passe hash\u00e9s',
]
for item in sec_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.4 Veille technologique \u2014 Vuln\u00e9rabilit\u00e9s de s\u00e9curit\u00e9', level=2)

doc.add_paragraph(
    'Sujet : Les attaques par injection dans les applications web modernes (OWASP Top 10 - 2021)\n\n'
    'L\u2019OWASP publie r\u00e9guli\u00e8rement un classement des vuln\u00e9rabilit\u00e9s web les plus critiques. '
    'En 2021, les injections (SQL, NoSQL, OS, LDAP) sont class\u00e9es en 3e position (A03:2021).\n\n'
    'Mesures impl\u00e9ment\u00e9es dans Vite & Gourmand :'
)

measures = [
    'Injection SQL : Prisma ORM utilise des requ\u00eates param\u00e9tr\u00e9es par d\u00e9faut. Aucune requ\u00eate SQL brute.',
    'Injection NoSQL : Mongoose valide les sch\u00e9mas avant insertion. Pipelines d\u2019agr\u00e9gation natifs.',
    'XSS : React \u00e9chappe automatiquement les variables. Pas d\u2019utilisation de dangerouslySetInnerHTML.',
    'CSRF : JWT dans les headers Authorization au lieu de cookies.',
    'Broken Authentication : politique de mot de passe forte, hachage bcrypt, tokens JWT avec expiration.',
]
for m in measures:
    doc.add_paragraph(m, style='List Bullet')

doc.add_paragraph('Source : OWASP Top 10 - 2021, https://owasp.org/Top10/')

doc.add_page_break()

# === PARTIE 3 ===
h = doc.add_heading('Partie 3 : Recherche', level=1)
for r in h.runs:
    r.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

doc.add_heading('3.1 Situation de recherche \u00e0 partir d\u2019un site anglophone', level=2)

doc.add_paragraph(
    'Lors de l\u2019int\u00e9gration de Prisma 7 avec le driver adapter pattern, j\u2019ai rencontr\u00e9 '
    'une erreur : "Error: \'url\' is not allowed in \'datasource\' block when using Prisma 7 '
    'with driver adapters". Le block datasource de Prisma ne pouvait plus contenir '
    'l\u2019attribut url directement.\n\n'
    'J\u2019ai recherch\u00e9 la solution sur la documentation officielle de Prisma en anglais '
    'et sur GitHub Issues.\n\n'
    'Source : https://www.prisma.io/docs/orm/overview/databases/neon#how-to-connect-using-prisma-client-and-a-driver-adapter'
)

doc.add_heading('3.2 Extrait du site anglophone et traduction', level=2)

p = doc.add_paragraph()
run = p.add_run('Extrait en anglais :')
run.bold = True

p = doc.add_paragraph()
run = p.add_run(
    '"When using Prisma ORM with a driver adapter, the connection to the database is not '
    'handled by Prisma\'s built-in engine. Instead, it is handled by the driver adapter. '
    'This means you should not include the url field in the datasource block of your Prisma '
    'schema. Instead, you should configure the connection in your application code using the '
    'driver adapter."'
)
run.italic = True

p = doc.add_paragraph()
run = p.add_run('Traduction en fran\u00e7ais :')
run.bold = True

p = doc.add_paragraph()
run = p.add_run(
    '\u00ab Lorsque vous utilisez Prisma ORM avec un adaptateur de driver, la connexion \u00e0 la '
    'base de donn\u00e9es n\u2019est pas g\u00e9r\u00e9e par le moteur int\u00e9gr\u00e9 de Prisma. Elle est g\u00e9r\u00e9e '
    'par l\u2019adaptateur de driver. Cela signifie que vous ne devez pas inclure le champ url '
    'dans le bloc datasource de votre sch\u00e9ma Prisma. \u00c0 la place, vous devez configurer la '
    'connexion dans votre code applicatif en utilisant l\u2019adaptateur de driver. \u00bb'
)
run.italic = True

doc.add_paragraph(
    'Application dans le projet : Cette recherche m\u2019a permis de configurer correctement '
    'Prisma 7 avec le driver adapter pattern pour Neon (PostgreSQL serverless). Le fichier '
    'prisma.config.ts g\u00e8re d\u00e9sormais la connexion via @prisma/adapter-pg au lieu du '
    'champ url dans le sch\u00e9ma Prisma.'
)

doc.add_page_break()

# === PARTIE 4 ===
h = doc.add_heading('Partie 4 : Informations compl\u00e9mentaires', level=1)
for r in h.runs:
    r.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

doc.add_heading('4.1 Autres ressources utilis\u00e9es', level=2)

table = doc.add_table(rows=10, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Ressource'
table.rows[0].cells[1].text = 'Usage'
for cell in table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True

resources = [
    ('Documentation Next.js (nextjs.org/docs)', 'Configuration App Router, routing, metadata API'),
    ('Documentation NestJS (docs.nestjs.com)', 'Modules, guards, pipes, decorators, JWT strategy'),
    ('Documentation Prisma (prisma.io/docs)', 'Schema, migrations, driver adapter pattern, seed'),
    ('Documentation MongoDB (mongodb.com/docs)', 'Aggregation pipelines, Atlas setup'),
    ('Documentation Tailwind CSS v4', 'Configuration CSS-first, custom theme, responsive'),
    ('Documentation Framer Motion (motion.dev)', 'API d\u2019animation, whileInView, variants'),
    ('MDN Web Docs (developer.mozilla.org)', 'R\u00e9f\u00e9rences HTML, CSS, JavaScript'),
    ('Stack Overflow', 'R\u00e9solution de bugs sp\u00e9cifiques'),
    ('OWASP (owasp.org)', 'Bonnes pratiques de s\u00e9curit\u00e9 web'),
]
for i, (res, usage) in enumerate(resources):
    table.rows[i+1].cells[0].text = res
    table.rows[i+1].cells[1].text = usage

doc.add_heading('4.2 Informations compl\u00e9mentaires', level=2)

doc.add_paragraph(
    'Architecture technique double BDD : Le projet utilise PostgreSQL pour les donn\u00e9es '
    'relationnelles (utilisateurs, menus, commandes) et MongoDB Atlas pour les statistiques '
    'agr\u00e9g\u00e9es du dashboard admin. Cette architecture r\u00e9pond \u00e0 l\u2019exigence ECF d\u2019utiliser '
    'deux types de bases de donn\u00e9es (SQL + NoSQL) tout en apportant une vraie valeur '
    'ajout\u00e9e : les pipelines d\u2019agr\u00e9gation MongoDB sont plus performants que les GROUP BY '
    'SQL pour les calculs statistiques en temps r\u00e9el.'
)

doc.add_paragraph(
    'Accessibilit\u00e9 : L\u2019application respecte les crit\u00e8res de base du RGAA : labels associ\u00e9s '
    '\u00e0 tous les champs de formulaire, contraste de couleurs suffisant, navigation au clavier, '
    'textes alternatifs sur les images, structure s\u00e9mantique HTML (header, main, footer, nav, section).'
)

doc.add_paragraph(
    'Performance : L\u2019application utilise les optimisations natives de Next.js (code splitting '
    'automatique, optimisation d\u2019images, polices auto-h\u00e9berg\u00e9es) et Framer Motion avec des '
    'animations GPU-acc\u00e9l\u00e9r\u00e9es (transform/opacity uniquement).'
)

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
    'ECF_TPDeveloppeurWebEtWebMobile_copiearendre_GAILLARD_Romain.docx')
doc.save(output_path)
size = os.path.getsize(output_path)
print(f'Word document saved: {output_path} ({size:,} bytes)')
